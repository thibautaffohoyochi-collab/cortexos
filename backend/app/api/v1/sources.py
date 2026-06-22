"""
CortexOS — Data Sources Routes
POST /sources/upload  → upload CSV, TXT, PDF, Excel or Word, ingest into Qdrant
GET  /sources         → list data sources for tenant
DELETE /sources/{id}  → delete a source
"""
import csv
import io
import uuid

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.core.database import get_db
from app.core.auth import get_current_user
from app.models.models import User, DataSource, SourceType, SourceStatus
from app.services.qdrant_service import upsert_chunks, chunk_text

router = APIRouter(prefix="/sources", tags=["sources"])

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB (raised for PDFs/Excel)

SUPPORTED_EXTENSIONS = ("csv", "txt", "pdf", "xlsx", "xls", "docx", "doc")


# ─── Parsers ──────────────────────────────────────────────────────────────────

def parse_csv(content: str, filename: str) -> list[dict]:
    """Parse CSV → chunk dicts."""
    reader = csv.DictReader(io.StringIO(content))
    rows = list(reader)
    if not rows:
        return []
    all_text = ""
    for row in rows:
        line = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
        all_text += line + "\n"
    chunks = chunk_text(all_text, chunk_size=300, overlap=30)
    return [{"text": c, "title": filename, "index": i} for i, c in enumerate(chunks)]


def parse_txt(content: str, filename: str) -> list[dict]:
    """Parse plain text → chunk dicts."""
    chunks = chunk_text(content, chunk_size=400, overlap=50)
    return [{"text": c, "title": filename, "index": i} for i, c in enumerate(chunks)]


def parse_pdf(raw: bytes, filename: str) -> list[dict]:
    """Parse PDF → chunk dicts using PyPDF2."""
    try:
        import PyPDF2
    except ImportError:
        raise HTTPException(status_code=500, detail="pypdf2 non installé sur le serveur.")

    reader = PyPDF2.PdfReader(io.BytesIO(raw))
    pages_text = []
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(f"[Page {page_num}]\n{text}")

    full_text = "\n\n".join(pages_text)
    if not full_text.strip():
        raise HTTPException(status_code=400, detail="Le PDF ne contient pas de texte extractible (PDF scanné ou protégé).")

    chunks = chunk_text(full_text, chunk_size=400, overlap=50)
    return [{"text": c, "title": f"{filename} (p.{i+1})", "index": i} for i, c in enumerate(chunks)]


def parse_excel(raw: bytes, filename: str) -> list[dict]:
    """Parse Excel (.xlsx/.xls) → chunk dicts using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl non installé sur le serveur.")

    wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    all_text = ""

    for sheet in wb.worksheets:
        all_text += f"\n## Feuille : {sheet.title}\n"
        rows_data = []
        headers: list[str] = []

        for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
            # Skip fully empty rows
            if all(cell is None for cell in row):
                continue
            row_values = [str(cell) if cell is not None else "" for cell in row]
            if row_idx == 0:
                headers = row_values
            else:
                if headers:
                    line = " | ".join(
                        f"{h}: {v}" for h, v in zip(headers, row_values) if v
                    )
                else:
                    line = " | ".join(v for v in row_values if v)
                if line.strip():
                    rows_data.append(line)

        all_text += "\n".join(rows_data) + "\n"

    if not all_text.strip():
        raise HTTPException(status_code=400, detail="Le fichier Excel est vide.")

    chunks = chunk_text(all_text, chunk_size=300, overlap=30)
    return [{"text": c, "title": filename, "index": i} for i, c in enumerate(chunks)]


def parse_word(raw: bytes, filename: str) -> list[dict]:
    """Parse Word (.docx) → chunk dicts using python-docx."""
    try:
        import docx
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx non installé sur le serveur.")

    doc = docx.Document(io.BytesIO(raw))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Keep headings with a marker
            if para.style.name.startswith("Heading"):
                paragraphs.append(f"\n## {text}")
            else:
                paragraphs.append(text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs)
    if not full_text.strip():
        raise HTTPException(status_code=400, detail="Le document Word est vide.")

    chunks = chunk_text(full_text, chunk_size=400, overlap=50)
    return [{"text": c, "title": filename, "index": i} for i, c in enumerate(chunks)]


# ─── Extension → parser + SourceType mapping ──────────────────────────────────

EXT_MAP = {
    "csv":  (parse_csv,   SourceType.CSV),
    "txt":  (parse_txt,   SourceType.CSV),       # reuse CSV type for plain text
    "pdf":  (parse_pdf,   SourceType.CSV),        # no PDF enum yet, use CSV as generic
    "xlsx": (parse_excel, SourceType.EXCEL),
    "xls":  (parse_excel, SourceType.EXCEL),
    "docx": (parse_word,  SourceType.CSV),        # no DOCX enum, use CSV as generic
    "doc":  (parse_word,  SourceType.CSV),
}


# ─── Routes ───────────────────────────────────────────────────────────────────

@router.post("/upload")
async def upload_source(
    file: UploadFile = File(...),
    name: str = Form(default=""),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    filename = file.filename or "document"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Format non supporté (.{ext}). Formats acceptés : {', '.join('.' + e for e in SUPPORTED_EXTENSIONS)}"
        )

    raw = await file.read()
    if len(raw) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="Fichier trop volumineux (max 20 MB).")

    if len(raw) == 0:
        raise HTTPException(status_code=400, detail="Le fichier est vide.")

    source_name = name or filename.rsplit(".", 1)[0]
    parser_fn, source_type = EXT_MAP[ext]

    # Parse into chunks
    try:
        if ext in ("csv", "txt"):
            content = raw.decode("utf-8", errors="replace")
            chunks = parser_fn(content, source_name)  # type: ignore
        else:
            chunks = parser_fn(raw, source_name)  # type: ignore
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erreur de parsing : {str(e)}")

    if not chunks:
        raise HTTPException(status_code=400, detail="Le fichier est vide ou non parseable.")

    # Create DataSource record
    source = DataSource(
        tenant_id=current_user.tenant_id,
        name=source_name,
        source_type=source_type,
        status=SourceStatus.SYNCING,
    )
    db.add(source)
    await db.flush()

    # Ingest into Qdrant
    try:
        count = await upsert_chunks(
            chunks=chunks,
            tenant_id=str(current_user.tenant_id),
            source_id=str(source.id),
        )
        source.status = SourceStatus.ACTIVE
        source.config = {"chunk_count": count, "filename": filename, "format": ext}
    except Exception as e:
        source.status = SourceStatus.ERROR
        source.error_message = str(e)
        import traceback
        raise HTTPException(status_code=500, detail=f"{type(e).__name__}: {e}\n{traceback.format_exc()}")

    return {
        "id": str(source.id),
        "name": source.name,
        "status": source.status,
        "chunk_count": len(chunks),
        "format": ext.upper(),
    }


@router.get("")
async def list_sources(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(DataSource)
        .where(DataSource.tenant_id == current_user.tenant_id)
        .order_by(DataSource.created_at.desc())
    )
    sources = result.scalars().all()
    return [
        {
            "id": str(s.id),
            "name": s.name,
            "source_type": s.source_type,
            "status": s.status,
            "created_at": s.created_at.isoformat(),
            "error_message": s.error_message,
            "format": (s.config or {}).get("format", "").upper() if s.config else "",
        }
        for s in sources
    ]


@router.delete("/{source_id}")
async def delete_source(
    source_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(DataSource).where(
            DataSource.id == source_id,
            DataSource.tenant_id == current_user.tenant_id,
        )
    )
    source = result.scalar_one_or_none()
    if not source:
        raise HTTPException(status_code=404, detail="Source introuvable")

    await db.delete(source)
    return {"ok": True}
